#! python3
# Generates an invitation for each name in a guests.txt file

# 1. Check the validity of the text file
#   a) check if the text file exists
#   b) check if it is not empty
# 2. Read the text file line by line
#   a) add each line to a list of names
# 3. Prepare the text of the document
#   a) create a list with custom text and 'name' to be replaced
# 4. Prepare the document itself
#   a) create a document object
#   b) for every name in name list:
#       add paragraph and style it
#       if there are names left add break

import sys, docx, os.path
from docx import Document
#from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. Check the validity of the text file
#   a) check if the text file exists
fpath = os.getcwd() + "/guests.txt"
if not os.path.isfile(fpath):
    print("guests.txt is not in the current working directory")
    sys.exit(1)

#   b) check if it is not empty
if not os.path.getsize(fpath):
    print("guests.txt is empty")
    sys.exit(1)

# 2. Read the text file line by line
#   a) add each line to a list of names
names = []
guestFile = open(fpath, 'r')
names = guestFile.readlines()
guestFile.close()

# 3. Prepare the text of the document
#   a) create a list with custom text and 'name' to be replaced
textLst = [["It would be a pleasure to have the company of"], ["name"], 
    ["at", " 11010 Memory Lane on the Evening of"], ["April 1st"], 
    ["at", " 7 o'clock"]]

# 4. Prepare the document itself
#   a) create a document object
doc = Document()
doc.save("invitations.docx")

#   b) for every name in name list:
for name_n in range(len(names)):
    print("Working with " + names[name_n].rstrip() + "...")
    #       add paragraph and style it
    for line_n in range(len(textLst)):
        print("Adding line " + str(line_n + 1) + "...")
        if textLst[line_n][0] == "name":
            paraObj = doc.add_paragraph(names[name_n].rstrip(), 'Normal')
            run = paraObj.runs[0]
            font = run.font
            font.size = docx.shared.Pt(18)
            font.bold = True
            #doc.paragraphs[len(doc.paragraphs)-1].runs[0].underline = True
        else:
            paraObj = doc.add_paragraph(textLst[line_n][0], 'Normal')
            if not line_n % 2:
                run = paraObj.runs[0]
                font = run.font
                font.size = docx.shared.Pt(20)
                font.name = "Brush Script Std"
                if line_n != 0:
                    font.underline = True
            else:
                run = paraObj.runs[0]
                font = run.font
                font.size = docx.shared.Pt(18)


        paraObj.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        if len(textLst[line_n]) > 1:
            for run_n in range(1, len(textLst[line_n])):
                paraObj.add_run(textLst[line_n][run_n])
                if not line_n % 2:
                    run = paraObj.runs[1]
                    font = run.font
                    font.size = docx.shared.Pt(20)
                    font.name = "Brush Script Std"

    #       if there are names left add break
    if name_n:
        doc.paragraphs[len(doc.paragraphs) - (len(textLst)+1)].runs[1].add_break(docx.enum.text.WD_BREAK.PAGE)

doc.save("invitations.docx")
print("Done")
